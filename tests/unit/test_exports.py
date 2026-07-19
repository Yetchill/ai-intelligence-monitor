# pyright: reportPrivateUsage=false, reportUnknownMemberType=false
# pyright: reportUnknownVariableType=false
"""Stage-six Excel/Word exports using only the temporary database fixture."""

from collections.abc import Iterator
from datetime import UTC, datetime, timedelta
from hashlib import sha256
from io import BytesIO
from pathlib import Path
from typing import cast
from xml.etree import ElementTree
from zipfile import ZipFile

import pytest
from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient
from openpyxl import load_workbook
from openpyxl.cell.cell import Cell
from openpyxl.worksheet.worksheet import Worksheet
from sqlalchemy import event, insert, select

from app import cli
from app.domain.enums import (
    Category,
    SourceAudience,
    SourceKind,
    SourceOrigin,
    SourceTier,
    SourceType,
)
from app.domain.exports import (
    EmptyExportError,
    ExportFormat,
    ExportLimitExceededError,
    ExportQuery,
    InvalidExportLimitError,
)
from app.domain.models import Base, IntelligenceItem, Source
from app.domain.queries import ItemFilter, ItemQuery
from app.exporters.common import CATEGORY_LABELS, CATEGORY_ORDER, excel_safe_text, safe_filename
from app.exporters.excel import ExcelExporter
from app.exporters.word import WordExporter
from app.services.application_factory import build_export_service
from app.services.web_data_service import WebDataService
from app.storage.database import Database
from app.storage.repositories import RepositoryUnitOfWork
from app.web.app import create_app
from app.web.routes.pages import _content_disposition
from app.web.schemas import WebInputError


@pytest.fixture
def export_app(database: Database) -> FastAPI:
    return create_app(database=database, enforce_migrations=False)


@pytest.fixture
def export_client(export_app: FastAPI) -> Iterator[TestClient]:
    with TestClient(export_app, raise_server_exceptions=False) as client:
        yield client


def _source(name: str, url: str) -> Source:
    return Source(
        name=name,
        source_type=SourceType.RSS,
        start_url=url,
        collector_name="rss",
        collector_config={},
        origin=SourceOrigin.PRESET,
        source_kind=SourceKind.FORMAL,
        source_tier=SourceTier.OFFICIAL_COMPANY,
        audience=SourceAudience.LEADERSHIP,
        homepage_visible=True,
        export_visible=True,
    )


def _seed(database: Database) -> tuple[Source, Source]:
    first = _source("第一来源", "https://first.example/feed")
    second = _source("第二来源", "https://second.example/feed")
    now = datetime(2026, 7, 18, 8, 0, tzinfo=UTC)
    with RepositoryUnitOfWork(database) as uow:
        uow.sources.add(first)
        uow.sources.add(second)
        rows = (
            (
                first,
                "人工奖项",
                "Agent 关键词简介",
                Category.MODEL_TECHNOLOGY,
                Category.AWARD_CASE,
                True,
                0,
            ),
            (second, "待分类无简介", None, Category.UNCLASSIFIED, None, False, 1),
            (first, "企业案例", "企业落地案例", Category.ENTERPRISE_CASE, None, False, 2),
            (second, "大模型发布", "模型能力更新", Category.MODEL_TECHNOLOGY, None, False, 3),
        )
        for source, title, summary, category, manual, favorite, days in rows:
            uow.items.add(
                IntelligenceItem(
                    source_id=source.id,
                    title=title,
                    summary=summary,
                    original_url=f"https://articles.example/{days}?a=1&b=2",
                    canonical_url=f"https://articles.example/{days}?a=1&b=2",
                    published_at=now - timedelta(days=days),
                    discovered_at=now - timedelta(hours=days),
                    last_seen_at=now,
                    category=category,
                    manual_category=manual,
                    classification_score=0.85 - days / 100,
                    classification_reason=f"规则原因 {title}",
                    fingerprint=f"{days:064x}",
                    is_favorite=favorite,
                )
            )
    return first, second


def _excel_result(database: Database, item_filter: ItemFilter | None = None):  # type: ignore[no-untyped-def]
    return build_export_service(database).export(
        ExportFormat.EXCEL,
        ExportQuery(filters=item_filter or ItemFilter()),
    )


def _excel_sheet(content: bytes) -> tuple[object, Worksheet]:
    workbook = load_workbook(BytesIO(content))
    return workbook, cast(Worksheet, workbook["资讯列表"])


def _excel_titles(content: bytes) -> list[str]:
    workbook, sheet = _excel_sheet(content)
    try:
        return [str(sheet.cell(row=row, column=2).value) for row in range(2, sheet.max_row + 1)]
    finally:
        workbook.close()  # type: ignore[attr-defined]


def _word_titles(content: bytes) -> list[str]:
    document = Document(BytesIO(content))
    return [
        paragraph.text.split(". ", maxsplit=1)[1]
        for paragraph in document.paragraphs
        if paragraph.style is not None and paragraph.style.name == "Heading 2"
    ]


def _bulk_seed(database: Database, count: int) -> None:
    source = _source("边界来源", "https://boundary.example/feed")
    now = datetime(2026, 7, 18, 8, 0, tzinfo=UTC)
    with database.session() as session:
        session.add(source)
        session.flush()
        session.execute(
            insert(IntelligenceItem),
            [
                {
                    "source_id": source.id,
                    "title": f"边界条目 {index}",
                    "summary": None,
                    "original_url": f"https://boundary.example/items/{index}",
                    "canonical_url": f"https://boundary.example/items/{index}",
                    "published_at": None,
                    "discovered_at": now,
                    "last_seen_at": now,
                    "updated_at": now,
                    "category": Category.MODEL_TECHNOLOGY,
                    "classification_score": None,
                    "classification_reason": None,
                    "automatic_category_provider": None,
                    "manual_category": None,
                    "fingerprint": f"{index:064x}",
                    "is_favorite": False,
                    "is_active": True,
                    "extra": {},
                }
                for index in range(count)
            ],
        )


def test_excel_normal_export_structure_and_safe_filename(database: Database) -> None:
    _seed(database)
    result = _excel_result(database)
    workbook, sheet = _excel_sheet(result.content)
    try:
        assert result.item_count == 4
        assert result.media_type.endswith("spreadsheetml.sheet")
        assert result.filename.endswith(".xlsx") and "/" not in result.filename
        assert workbook.sheetnames == ["资讯列表", "导出说明"]  # type: ignore[attr-defined]
        assert sheet.freeze_panes == "A2"
        assert sheet.auto_filter.ref == "A1:M5"
        header_cells = [cast(Cell, cell) for cell in sheet[1]]
        assert [cell.value for cell in header_cells] == list(ExcelExporter._headers)
        assert all(cast(bool, cell.font.bold) for cell in header_cells)
        assert sheet.column_dimensions["B"].width == 44
        assert sheet.cell(row=2, column=9).hyperlink is not None
        assert sheet.cell(row=2, column=9).hyperlink.target.startswith("https://")  # type: ignore[union-attr]
        assert sheet.__dict__.get("_images") == []
    finally:
        workbook.close()  # type: ignore[attr-defined]


def test_manual_category_priority_and_chinese_category(database: Database) -> None:
    _seed(database)
    workbook, sheet = _excel_sheet(_excel_result(database).content)
    try:
        rows = {sheet.cell(row=row, column=2).value: row for row in range(2, sheet.max_row + 1)}
        manual_row = rows["人工奖项"]
        assert sheet.cell(row=manual_row, column=3).value == CATEGORY_LABELS[Category.AWARD_CASE]
        assert sheet.cell(row=manual_row, column=4).value == "人工"
    finally:
        workbook.close()  # type: ignore[attr-defined]


@pytest.mark.parametrize(
    ("item_filter", "expected"),
    [
        (ItemFilter(category=Category.AWARD_CASE), ["人工奖项"]),
        (ItemFilter(favorite=True), ["人工奖项"]),
        (ItemFilter(keyword="Agent"), ["人工奖项"]),
        (ItemFilter(unclassified=True), ["待分类无简介"]),
        (
            ItemFilter(
                published_from=datetime(2026, 7, 17, tzinfo=UTC),
                published_to=datetime(2026, 7, 18, tzinfo=UTC),
            ),
            ["待分类无简介"],
        ),
        (
            ItemFilter(
                discovered_from=datetime(2026, 7, 18, 6, tzinfo=UTC),
                discovered_to=datetime(2026, 7, 18, 7, tzinfo=UTC),
            ),
            ["企业案例"],
        ),
    ],
)
def test_export_filters(
    database: Database,
    item_filter: ItemFilter,
    expected: list[str],
) -> None:
    _seed(database)
    assert _excel_titles(_excel_result(database, item_filter).content) == expected


def test_source_filter_and_multi_condition_and(database: Database) -> None:
    first, _ = _seed(database)
    item_filter = ItemFilter(
        source_id=first.id,
        favorite=False,
        keyword="案例",
        category=Category.ENTERPRISE_CASE,
    )
    assert _excel_titles(_excel_result(database, item_filter).content) == ["企业案例"]


@pytest.mark.parametrize("prefix", ["=SUM(1,1)", "+cmd", "+42", "-1+1", "@example"])
def test_excel_formula_injection_is_plain_text(database: Database, prefix: str) -> None:
    source = _source(prefix, f"https://formula-{len(prefix)}.example/feed")
    with RepositoryUnitOfWork(database) as uow:
        uow.sources.add(source)
        uow.items.add(
            IntelligenceItem(
                source_id=source.id,
                title=prefix,
                summary=prefix,
                original_url="https://formula.example/item",
                canonical_url="https://formula.example/item",
                category=Category.MODEL_TECHNOLOGY,
                classification_reason=prefix,
                fingerprint="f" * 64,
            )
        )
    workbook, sheet = _excel_sheet(_excel_result(database).content)
    try:
        for column in (2, 5, 8, 12):
            cell = sheet.cell(row=2, column=column)
            assert cell.data_type != "f"
            assert str(cell.value).startswith("'")
    finally:
        workbook.close()  # type: ignore[attr-defined]


@pytest.mark.parametrize(
    "value",
    [
        " =SUM(1,1)",
        "\t =SUM(1,1)",
        "\r\n@cmd",
        "\x00\x7f-1+1",
        "\ufeff+cmd",
    ],
)
def test_excel_formula_injection_after_whitespace_and_controls_is_plain_text(
    value: str,
) -> None:
    rendered = excel_safe_text(value)
    assert rendered.startswith("'")
    assert len(rendered) <= 32_767


@pytest.mark.parametrize(
    "value",
    ["-123", "-12.50", "-1e3", "2026-07-18", "普通文本", "说明 - 正常内容"],
)
def test_excel_formula_protection_preserves_normal_values(value: str) -> None:
    assert excel_safe_text(value) == value


def test_excel_formula_protection_respects_string_length_limit() -> None:
    rendered = excel_safe_text("=" + "长" * 40_000)
    assert rendered.startswith("'=")
    assert len(rendered) == 32_767


def test_malicious_long_unicode_is_safely_written(database: Database) -> None:
    source = _source("<script>来源\x00", "https://malicious.example/feed")
    long_summary = "<img src=x onerror=alert(1)>\ufffe" + "长" * 40_000
    with RepositoryUnitOfWork(database) as uow:
        uow.sources.add(source)
        uow.items.add(
            IntelligenceItem(
                source_id=source.id,
                title="<script>alert(1)</script>\x00",
                summary=long_summary,
                original_url="https://malicious.example/item",
                canonical_url="https://malicious.example/item",
                fingerprint="a" * 64,
            )
        )
    excel = _excel_result(database)
    workbook, sheet = _excel_sheet(excel.content)
    try:
        assert "\x00" not in str(sheet.cell(2, 2).value)
        assert len(str(sheet.cell(2, 8).value)) <= 32_767
        assert "<script>" in str(sheet.cell(2, 2).value)
    finally:
        workbook.close()  # type: ignore[attr-defined]
    word = build_export_service(database).export(ExportFormat.WORD, ExportQuery())
    document = Document(BytesIO(word.content))
    assert "<script>" in "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert len(document.inline_shapes) == 0


def test_office_packages_reopen_and_have_only_expected_parts(database: Database) -> None:
    _seed(database)
    results = (
        _excel_result(database),
        build_export_service(database).export(ExportFormat.WORD, ExportQuery()),
    )
    for result in results:
        with ZipFile(BytesIO(result.content)) as package:
            assert package.testzip() is None
            names = set(package.namelist())
            assert not any(
                forbidden in name.lower()
                for name in names
                for forbidden in ("vbaproject", "externallinks", "embeddings/", "media/")
            )
            for name in names:
                if name.endswith((".xml", ".rels")):
                    ElementTree.fromstring(package.read(name))
            if result.media_type == WordExporter.media_type:
                styles = package.read("word/styles.xml")
                assert b"eastAsia" in styles
    workbook = load_workbook(BytesIO(results[0].content))
    workbook.close()
    Document(BytesIO(results[1].content))


def test_word_structure_group_order_no_empty_sections_and_hyperlinks(database: Database) -> None:
    _seed(database)
    result = build_export_service(database).export(ExportFormat.WORD, ExportQuery())
    document = Document(BytesIO(result.content))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert result.item_count == 4
    assert result.media_type.endswith("wordprocessingml.document")
    assert "AI行业动态与成果申报情报汇总" in text
    assert "条目总数: 4" in text
    assert text.index("大模型与技术") < text.index("企业成果与案例")
    assert text.index("企业成果与案例") < text.index("奖项与优秀案例")
    assert text.index("奖项与优秀案例") < text.index("待分类")
    assert "智能体与产品" not in text
    assert "征集与申报" not in text
    assert "政策、标准与行业" not in text
    assert text.count("简介: ") == 3
    hyperlinks = [
        relationship
        for relationship in document.part.rels.values()
        if relationship.reltype.endswith("/hyperlink")
    ]
    assert len(hyperlinks) == 4
    assert all(relationship.target_ref.startswith("https://") for relationship in hyperlinks)
    assert len(document.inline_shapes) == 0


@pytest.mark.parametrize("export_format", list(ExportFormat))
def test_empty_result_and_format_limits(database: Database, export_format: ExportFormat) -> None:
    service = build_export_service(database)
    with pytest.raises(EmptyExportError):
        service.export(export_format, ExportQuery())

    _seed(database)
    with pytest.raises(ExportLimitExceededError, match="超过"):
        service.export(export_format, ExportQuery(limit=2))
    maximum = (
        ExcelExporter.max_items if export_format is ExportFormat.EXCEL else WordExporter.max_items
    )
    with pytest.raises(InvalidExportLimitError):
        service.export(export_format, ExportQuery(limit=maximum + 1))


def test_export_does_not_mutate_business_data(database: Database) -> None:
    _seed(database)
    with RepositoryUnitOfWork(database) as uow:
        before = [
            (item.id, item.updated_at, item.is_favorite, item.manual_category)
            for item in uow.items.list()
        ]
        assert uow.crawl_runs.list() == [] and uow.revisions.list() == []

    service = build_export_service(database)
    service.export(ExportFormat.EXCEL, ExportQuery())
    service.export(ExportFormat.WORD, ExportQuery())

    with RepositoryUnitOfWork(database) as uow:
        after = [
            (item.id, item.updated_at, item.is_favorite, item.manual_category)
            for item in uow.items.list()
        ]
        assert before == after
        assert uow.crawl_runs.list() == [] and uow.revisions.list() == []


def test_web_exports_inherit_filters_ignore_page_and_set_safe_headers(
    database: Database,
    export_client: TestClient,
) -> None:
    _seed(database)
    response = export_client.post(
        "/exports/excel",
        data={"category": "model_technology", "page": "99", "per_page": "20"},
    )

    assert response.status_code == 200
    assert response.headers["content-type"].startswith("application/vnd.openxmlformats")
    disposition = response.headers["content-disposition"]
    assert 'filename="ai-intelligence-' in disposition
    assert "filename*=UTF-8''AI%E8%A1%8C%E4%B8%9A%E6%83%85%E6%8A%A5_" in disposition
    assert "\r" not in disposition and "\n" not in disposition
    assert _excel_titles(cast(bytes, response.content)) == ["大模型发布"]


def test_web_empty_and_invalid_filters_are_clear_errors(export_client: TestClient) -> None:
    empty = export_client.post("/exports/word")
    invalid = export_client.post("/exports/excel", data={"published_from": "2026-99-99"})
    duplicate = export_client.post(
        "/exports/excel",
        data={"category": ["award_case", "unclassified"]},
    )

    assert empty.status_code == 400 and "当前筛选没有结果" in empty.text
    assert invalid.status_code == 400 and "日期格式应为 YYYY-MM-DD" in invalid.text
    assert duplicate.status_code == 400 and "导出筛选参数无效" in duplicate.text
    assert all("Traceback" not in response.text for response in (empty, invalid, duplicate))


def test_home_shows_export_scope_and_inherits_current_filters(
    database: Database, export_client: TestClient
) -> None:
    _seed(database)
    response = export_client.get("/?category=award_case&favorite=yes")

    assert response.status_code == 200
    assert "导出当前筛选结果" in response.text
    assert "导出包含全部匹配记录" in response.text
    assert 'action="/exports/excel"' in response.text
    assert 'action="/exports/word"' in response.text
    assert 'name="category" value="award_case"' in response.text
    assert 'name="favorite" value="yes"' in response.text


def test_web_word_content_type_and_openable_document(
    database: Database, export_client: TestClient
) -> None:
    _seed(database)
    response = export_client.post("/exports/word", data={"unclassified": "yes"})

    assert response.status_code == 200
    assert response.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert 'filename="ai-intelligence-' in response.headers["content-disposition"]
    assert response.headers["content-disposition"].endswith(".docx")
    document = Document(BytesIO(cast(bytes, response.content)))
    assert "待分类无简介" in "\n".join(paragraph.text for paragraph in document.paragraphs)


def test_web_export_is_post_only_and_content_disposition_is_strict(
    export_client: TestClient,
) -> None:
    assert export_client.get("/exports/excel").status_code == 405
    header = _content_disposition("中文 报告.xlsx", "report.xlsx")
    assert header == (
        'attachment; filename="report.xlsx"; '
        "filename*=UTF-8''%E4%B8%AD%E6%96%87%20%E6%8A%A5%E5%91%8A.xlsx"
    )
    for filename, ascii_filename in (
        ("恶意\r\nX-Evil: yes.xlsx", "report.xlsx"),
        ("正常.xlsx", "报告.xlsx"),
        ("正常.xlsx", 'report.xlsx"; X-Evil=yes'),
    ):
        with pytest.raises(WebInputError):
            _content_disposition(filename, ascii_filename)


def test_web_generation_failure_returns_no_partial_download(
    database: Database,
    export_client: TestClient,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    _seed(database)

    def fail_render(*_args: object, **_kwargs: object) -> object:
        raise RuntimeError("/private/tmp/secret-office-part")

    monkeypatch.setattr(ExcelExporter, "render", fail_render)
    response = export_client.post("/exports/excel")
    assert response.status_code == 500
    assert response.headers["content-type"].startswith("text/html")
    assert "content-disposition" not in response.headers
    assert "secret-office-part" not in response.text


def test_export_query_order_matches_home_stable_order(database: Database) -> None:
    _seed(database)
    assert _excel_titles(_excel_result(database).content) == [
        "人工奖项",
        "待分类无简介",
        "企业案例",
        "大模型发布",
    ]


def test_null_published_at_order_is_stable_and_matches_home(database: Database) -> None:
    source = _source("无发布时间来源", "https://no-date.example/feed")
    fallback_time = datetime(2026, 7, 18, 8, 0, tzinfo=UTC)
    with RepositoryUnitOfWork(database) as uow:
        uow.sources.add(source)
        for index in range(3):
            uow.items.add(
                IntelligenceItem(
                    source_id=source.id,
                    title=f"无发布时间 {index}",
                    original_url=f"https://no-date.example/{index}",
                    canonical_url=f"https://no-date.example/{index}",
                    published_at=None,
                    discovered_at=fallback_time,
                    fingerprint=f"{index + 10:064x}",
                )
            )
    page = WebDataService(lambda: RepositoryUnitOfWork(database)).list_items(ItemQuery())
    home_titles = [item.title for item in page.entries]
    assert home_titles == ["无发布时间 2", "无发布时间 1", "无发布时间 0"]
    assert _excel_titles(_excel_result(database).content) == home_titles


def test_excel_and_word_contain_same_items_with_documented_group_order(
    database: Database,
) -> None:
    _seed(database)
    excel = _excel_result(database)
    word = build_export_service(database).export(ExportFormat.WORD, ExportQuery())
    workbook, sheet = _excel_sheet(excel.content)
    try:
        excel_rows = [
            (str(sheet.cell(row, 2).value), str(sheet.cell(row, 3).value))
            for row in range(2, sheet.max_row + 1)
        ]
    finally:
        workbook.close()  # type: ignore[attr-defined]
    expected_word_order = [
        title
        for category in CATEGORY_ORDER
        for title, label in excel_rows
        if label == CATEGORY_LABELS[category]
    ]
    assert _word_titles(word.content) == expected_word_order
    assert {title for title, _label in excel_rows} == set(_word_titles(word.content))


@pytest.mark.parametrize("keyword", ["%", "_", "\\", "'", '"', "中文"])
def test_export_search_treats_special_characters_safely_and_literally(
    database: Database,
    keyword: str,
) -> None:
    source = _source("搜索来源", "https://search-export.example/feed")
    with RepositoryUnitOfWork(database) as uow:
        uow.sources.add(source)
        uow.items.add(
            IntelligenceItem(
                source_id=source.id,
                title=f"包含字面搜索值 {keyword}",
                original_url="https://search-export.example/match",
                canonical_url="https://search-export.example/match",
                fingerprint="c" * 64,
            )
        )
        uow.items.add(
            IntelligenceItem(
                source_id=source.id,
                title="不相关标题",
                original_url="https://search-export.example/other",
                canonical_url="https://search-export.example/other",
                fingerprint="d" * 64,
            )
        )
    assert _excel_titles(_excel_result(database, ItemFilter(keyword=keyword)).content) == [
        f"包含字面搜索值 {keyword}"
    ]


def test_word_category_filter_has_only_selected_nonempty_chapter(database: Database) -> None:
    _seed(database)
    result = build_export_service(database).export(
        ExportFormat.WORD,
        ExportQuery(filters=ItemFilter(category=Category.ENTERPRISE_CASE)),
    )
    text = "\n".join(paragraph.text for paragraph in Document(BytesIO(result.content)).paragraphs)
    assert "企业成果与案例" in text and "企业案例" in text
    assert "大模型与技术" not in text and "待分类" not in text


def test_unsafe_original_url_is_not_written_as_hyperlink(database: Database) -> None:
    source = _source("来源", "https://unsafe-link.example/feed")
    with RepositoryUnitOfWork(database) as uow:
        uow.sources.add(source)
        uow.items.add(
            IntelligenceItem(
                source_id=source.id,
                title="不安全链接",
                original_url="javascript:alert(1)",
                canonical_url="https://unsafe-link.example/item",
                fingerprint="b" * 64,
            )
        )
    excel = _excel_result(database)
    workbook, sheet = _excel_sheet(excel.content)
    try:
        assert sheet.cell(2, 9).value == "链接不可用"
        assert sheet.cell(2, 9).hyperlink is None
    finally:
        workbook.close()  # type: ignore[attr-defined]
    word = build_export_service(database).export(ExportFormat.WORD, ExportQuery())
    document = Document(BytesIO(word.content))
    assert not any(
        relationship.reltype.endswith("/hyperlink") for relationship in document.part.rels.values()
    )


def test_export_uses_one_limit_plus_one_select(database: Database) -> None:
    _seed(database)
    selects: list[tuple[str, object]] = []

    def record_select(
        _connection: object,
        _cursor: object,
        statement: str,
        parameters: object,
        _context: object,
        _executemany: object,
    ) -> None:
        if statement.lstrip().upper().startswith("SELECT"):
            selects.append((statement, parameters))

    event.listen(database.engine, "before_cursor_execute", record_select)
    try:
        build_export_service(database).export(
            ExportFormat.EXCEL,
            ExportQuery(limit=4),
        )
    finally:
        event.remove(database.engine, "before_cursor_execute", record_select)
    assert len(selects) == 1
    assert "LIMIT" in selects[0][0].upper()
    assert 5 in cast(tuple[object, ...], selects[0][1])


@pytest.mark.parametrize(
    ("export_format", "maximum"),
    [(ExportFormat.EXCEL, 10_000), (ExportFormat.WORD, 2_000)],
)
def test_format_hard_limit_boundary_succeeds(
    database: Database,
    export_format: ExportFormat,
    maximum: int,
) -> None:
    _bulk_seed(database, maximum)
    result = build_export_service(database).export(export_format, ExportQuery())
    assert result.item_count == maximum
    if export_format is ExportFormat.EXCEL:
        workbook, sheet = _excel_sheet(result.content)
        try:
            assert sheet.max_row == maximum + 1
        finally:
            workbook.close()  # type: ignore[attr-defined]
    else:
        assert len(_word_titles(result.content)) == maximum


@pytest.mark.parametrize(
    ("export_format", "count", "query_limit"),
    [(ExportFormat.EXCEL, 10_001, 10_001), (ExportFormat.WORD, 2_001, 2_001)],
)
def test_format_hard_limit_plus_one_is_rejected_by_one_bounded_query(
    database: Database,
    export_format: ExportFormat,
    count: int,
    query_limit: int,
) -> None:
    _bulk_seed(database, count)
    selects: list[tuple[str, object]] = []

    def record_select(
        _connection: object,
        _cursor: object,
        statement: str,
        parameters: object,
        _context: object,
        _executemany: object,
    ) -> None:
        if statement.lstrip().upper().startswith("SELECT"):
            selects.append((statement, parameters))

    event.listen(database.engine, "before_cursor_execute", record_select)
    try:
        with pytest.raises(ExportLimitExceededError):
            build_export_service(database).export(export_format, ExportQuery())
    finally:
        event.remove(database.engine, "before_cursor_execute", record_select)
    assert len(selects) == 1
    assert "LIMIT" in selects[0][0].upper()
    assert query_limit in cast(tuple[object, ...], selects[0][1])


def test_cli_invalid_date_returns_nonzero_without_output(
    database: Database,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    capsys: pytest.CaptureFixture[str],
) -> None:
    _seed(database)

    def fake_from_settings(_cls: type[Database]) -> Database:
        return database

    monkeypatch.setattr(cli, "configure_logging", lambda: None)
    monkeypatch.setattr(cli.Database, "from_settings", classmethod(fake_from_settings))
    output = tmp_path / "invalid.xlsx"
    code = cli.main(["export", "excel", "--output", str(output), "--published-from", "not-a-date"])
    assert code == 2 and not output.exists()
    assert "expected YYYY-MM-DD" in capsys.readouterr().err


def test_cli_rejects_mismatched_output_extension(
    database: Database,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    _seed(database)

    def fake_from_settings(_cls: type[Database]) -> Database:
        return database

    monkeypatch.setattr(cli, "configure_logging", lambda: None)
    monkeypatch.setattr(cli.Database, "from_settings", classmethod(fake_from_settings))
    output = tmp_path / "wrong.docx"
    assert cli.main(["export", "excel", "--output", str(output)]) == 2
    assert not output.exists()


def test_cli_output_no_overwrite_force_and_query_consistency(
    database: Database,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    capsys: pytest.CaptureFixture[str],
) -> None:
    _seed(database)
    output = tmp_path / "nested" / "report.xlsx"

    def fake_from_settings(_cls: type[Database]) -> Database:
        return database

    monkeypatch.setattr(cli, "configure_logging", lambda: None)
    monkeypatch.setattr(cli.Database, "from_settings", classmethod(fake_from_settings))
    arguments = [
        "export",
        "excel",
        "--output",
        str(output),
        "--category",
        "award_case",
    ]
    assert cli.main(arguments) == 0
    assert output.exists()
    expected = _excel_titles(
        _excel_result(database, ItemFilter(category=Category.AWARD_CASE)).content
    )
    assert _excel_titles(output.read_bytes()) == expected

    original = output.read_bytes()
    assert cli.main(arguments) == 2
    assert output.read_bytes() == original
    assert "File exists" in capsys.readouterr().err
    assert cli.main([*arguments, "--force"]) == 0
    assert _excel_titles(output.read_bytes()) == expected


def test_cli_default_output_directory_and_temp_cleanup(
    database: Database,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    _seed(database)

    def fake_from_settings(_cls: type[Database]) -> Database:
        return database

    monkeypatch.setattr(cli, "configure_logging", lambda: None)
    monkeypatch.setattr(cli.Database, "from_settings", classmethod(fake_from_settings))
    monkeypatch.setattr(cli, "PROJECT_ROOT", tmp_path)
    assert cli.main(["export", "word"]) == 0
    outputs = list((tmp_path / "output").glob("*.docx"))
    assert len(outputs) == 1

    def fail_link(_source: Path, _target: Path) -> None:
        raise OSError("simulated link failure")

    failed_output = tmp_path / "failed.xlsx"
    monkeypatch.setattr(cli.os, "link", fail_link)
    with pytest.raises(OSError):
        cli._atomic_write(failed_output, b"content", force=False)
    assert not failed_output.exists()
    assert list(tmp_path.glob(".failed.xlsx.*.tmp")) == []


def test_cli_rejects_directory_and_symlink_targets_and_cleans_force_failure(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    directory_target = tmp_path / "directory.xlsx"
    directory_target.mkdir()
    with pytest.raises(ValueError, match="directory"):
        cli._atomic_write(directory_target, b"content", force=True)

    real_target = tmp_path / "real.xlsx"
    real_target.write_bytes(b"original")
    symlink_target = tmp_path / "link.xlsx"
    symlink_target.symlink_to(real_target)
    with pytest.raises(ValueError, match="symbolic link"):
        cli._atomic_write(symlink_target, b"replacement", force=True)
    assert symlink_target.is_symlink()
    assert real_target.read_bytes() == b"original"

    output = tmp_path / "replace.xlsx"
    output.write_bytes(b"original")

    def fail_replace(_source: Path, _target: Path) -> None:
        raise OSError("simulated replace failure")

    monkeypatch.setattr(cli.os, "replace", fail_replace)
    with pytest.raises(OSError, match="replace failure"):
        cli._atomic_write(output, b"replacement", force=True)
    assert output.read_bytes() == b"original"
    assert list(tmp_path.glob(".replace.xlsx.*.tmp")) == []


def test_export_preserves_all_database_rows_and_file_hash(database: Database) -> None:
    _seed(database)

    def snapshot() -> dict[str, list[dict[str, object]]]:
        with database.engine.connect() as connection:
            return {
                table.name: [
                    dict(row)
                    for row in connection.execute(select(table).order_by(table.c.id)).mappings()
                ]
                for table in Base.metadata.sorted_tables
            }

    database_path = Path(str(database.engine.url.database))
    before_rows = snapshot()
    before_hash = sha256(database_path.read_bytes()).hexdigest()
    service = build_export_service(database)
    service.export(ExportFormat.EXCEL, ExportQuery())
    service.export(ExportFormat.WORD, ExportQuery())
    assert snapshot() == before_rows
    assert sha256(database_path.read_bytes()).hexdigest() == before_hash


def test_output_is_gitignored_but_gitkeep_is_tracked_configuration() -> None:
    ignore_text = Path(".gitignore").read_text(encoding="utf-8")
    assert "output/*" in ignore_text
    assert "!output/.gitkeep" in ignore_text
    assert Path("output/.gitkeep").exists()


@pytest.mark.parametrize(
    ("unsafe", "expected"),
    [
        ("../../报告:*?", "报告.xlsx"),
        ("..", "export.xlsx"),
        ("正常文件", "正常文件.xlsx"),
    ],
)
def test_safe_filename(unsafe: str, expected: str) -> None:
    assert safe_filename(unsafe, extension="xlsx") == expected
