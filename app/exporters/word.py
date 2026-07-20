# pyright: reportUnknownMemberType=false, reportUnknownVariableType=false
# pyright: reportPrivateUsage=false
"""Plain-text Word report grouped by effective category."""

from collections.abc import Sequence
from io import BytesIO
from typing import cast

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.styles.style import ParagraphStyle
from docx.text.paragraph import Paragraph

from app.domain.exports import ExportFormat, ExportItem, ExportMetadata, RenderedExport
from app.exporters.common import (
    PRIMARY_TYPE_LABELS,
    PRIMARY_TYPE_ORDER,
    clean_text,
    format_time,
    safe_filename,
    safe_hyperlink,
)


class WordExporter:
    export_format = ExportFormat.WORD
    max_items = 2_000
    media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def render(
        self,
        items: Sequence[ExportItem],
        metadata: ExportMetadata,
    ) -> RenderedExport:
        document = Document()
        buffer = BytesIO()
        try:
            _configure_document(document)
            title = document.add_heading("AI行业动态与成果申报情报汇总", level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            intro = document.add_paragraph()
            intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
            intro.add_run(f"生成时间: {format_time(metadata.generated_at)}\n")
            intro.add_run(f"筛选条件: {clean_text(metadata.filter_summary, limit=2_000)}\n")
            intro.add_run(f"条目总数: {metadata.item_count}")

            grouped = {
                primary_type: [
                    item for item in items if item.effective_primary_type is primary_type
                ]
                for primary_type in PRIMARY_TYPE_ORDER
            }
            chapter_number = 0
            for primary_type in PRIMARY_TYPE_ORDER:
                type_items = grouped[primary_type]
                if not type_items:
                    continue
                chapter_number += 1
                heading = document.add_heading(
                    f"{_chinese_number(chapter_number)}、{PRIMARY_TYPE_LABELS[primary_type]}",
                    level=1,
                )
                if chapter_number > 1:
                    heading.paragraph_format.page_break_before = True
                for item_number, item in enumerate(type_items, start=1):
                    document.add_heading(
                        f"{item_number}. {clean_text(item.title, limit=1_000)}", level=2
                    )
                    details = document.add_paragraph()
                    details.add_run("来源: ").bold = True
                    details.add_run(
                        f"{clean_text(item.source_name, limit=255)} ({item.source_kind.value})"
                    )
                    details.add_run("\n发布时间: ").bold = True
                    details.add_run(format_time(item.published_at) or "未知")
                    details.add_run("\n分类方式: ").bold = True
                    details.add_run(item.category_origin)
                    details.add_run("\n可信状态: ").bold = True
                    details.add_run(item.verification_status.value)
                    details.add_run("\n审核状态: ").bold = True
                    details.add_run(item.review_status.value)
                    if item.topic_tags:
                        details.add_run("\n主题标签: ").bold = True
                        details.add_run(", ".join(item.topic_tags))
                    if item.industry_tags:
                        details.add_run("\n行业标签: ").bold = True
                        details.add_run(", ".join(item.industry_tags))
                    if item.summary:
                        summary = document.add_paragraph()
                        summary.add_run("简介: ").bold = True
                        summary.add_run(clean_text(item.summary, limit=20_000))
                    link = safe_hyperlink(item.original_url)
                    if link is not None:
                        link_paragraph = document.add_paragraph()
                        link_paragraph.add_run("原文: ").bold = True
                        _add_hyperlink(link_paragraph, link, "查看原文")
                    official_link = safe_hyperlink(item.official_url)
                    if official_link is not None and official_link != link:
                        official_paragraph = document.add_paragraph()
                        official_paragraph.add_run("官方原文: ").bold = True
                        _add_hyperlink(official_paragraph, official_link, "查看官方原文")

            document.save(buffer)
            date_stamp = metadata.generated_at.strftime("%Y%m%d")
            return RenderedExport(
                content=buffer.getvalue(),
                filename=safe_filename(f"AI行业情报_{date_stamp}", extension="docx"),
                ascii_filename=safe_filename(f"ai-intelligence-{date_stamp}", extension="docx"),
                media_type=self.media_type,
            )
        finally:
            buffer.close()


def _configure_document(document: DocumentObject) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    for style_name, size, bold in (
        ("Normal", 10.5, False),
        ("Title", 22, True),
        ("Heading 1", 16, True),
        ("Heading 2", 12, True),
    ):
        style = cast(ParagraphStyle, document.styles[style_name])
        style.font.name = "Microsoft YaHei"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.25
    heading_two = cast(ParagraphStyle, document.styles["Heading 2"])
    heading_two.paragraph_format.keep_with_next = True


def _add_hyperlink(paragraph: Paragraph, url: str, label: str) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), str(RGBColor(5, 99, 193)))
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(color)
    properties.append(underline)
    run.append(properties)
    text = OxmlElement("w:t")
    text.text = clean_text(label, limit=100)
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _chinese_number(value: int) -> str:
    return ("一", "二", "三", "四", "五", "六", "七", "八")[value - 1]
